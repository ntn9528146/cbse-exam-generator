import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_number_field(run):
    """Inserts a dynamic Word XML Page Number field into a footer run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def set_cell_margins_zero(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_and_save_docx(school_name, class_level, subject, total_marks, time_allowed, syllabus, sections_list, general_instructions, logo_path=None):
    doc = Document()

    # 1. Global Document Styles (Times New Roman, 12pt Normal)
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_font.color.rgb = RGBColor(0, 0, 0)

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # ---------------- Header Configuration (Every Page) ----------------
        header = section.header
        header.is_linked_to_previous = False
        
        header_table = header.add_table(rows=1, cols=2, width=Inches(7.0))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.3)
        header_table.columns[1].width = Inches(5.7)

        # Left: School Logo
        cell_logo = header_table.cell(0, 0)
        set_cell_margins_zero(cell_logo)
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if logo_path and os.path.exists(logo_path):
            try:
                p_logo.add_run().add_picture(logo_path, width=Inches(0.95))
            except Exception:
                pass

        # Center: School Name (14pt, Bold, Times New Roman)
        cell_name = header_table.cell(0, 1)
        set_cell_margins_zero(cell_name)
        p_name = cell_name.paragraphs[0]
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_name = p_name.add_run(school_name.upper())
        r_name.font.name = 'Times New Roman'
        r_name.font.size = Pt(14)
        r_name.bold = True

        p_subhead = cell_name.add_paragraph()
        p_subhead.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sh = p_subhead.add_run(f"EXAMINATION (SESSION 2026-2027) | {class_level.upper()}")
        r_sh.font.name = 'Times New Roman'
        r_sh.font.size = Pt(11)
        r_sh.bold = True

        # ---------------- Footer Configuration (Page Number) ----------------
        footer = section.footer
        footer.is_linked_to_previous = False
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_foot_label = p_foot.add_run("Page ")
        r_foot_label.font.name = 'Times New Roman'
        r_foot_label.font.size = Pt(10)
        
        r_foot_pg = p_foot.add_run()
        r_foot_pg.font.name = 'Times New Roman'
        r_foot_pg.font.size = Pt(10)
        add_page_number_field(r_foot_pg)

    # 2. Main Page Header Info Table (Roll No, Marks, Time)
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main_sub = p_main_title.add_run(f"SUBJECT: {subject.upper()}\n")
    r_main_sub.font.name = 'Times New Roman'
    r_main_sub.font.size = Pt(14)
    r_main_sub.bold = True

    info_table = doc.add_table(rows=1, cols=2)
    info_table.autofit = False
    info_table.columns[0].width = Inches(3.5)
    info_table.columns[1].width = Inches(3.5)
    
    p_time = info_table.cell(0, 0).paragraphs[0]
    r_t = p_time.add_run(f"Time Allowed: {time_allowed}")
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(12)
    r_t.bold = True

    p_mm = info_table.cell(0, 1).paragraphs[0]
    p_mm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_m = p_mm.add_run(f"Maximum Marks: {total_marks}")
    r_m.font.name = 'Times New Roman'
    r_m.font.size = Pt(12)
    r_m.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # General Instructions
    inst_heading = doc.add_paragraph()
    r_ih = inst_heading.add_run("General Instructions:")
    r_ih.font.name = 'Times New Roman'
    r_ih.font.size = Pt(12)
    r_ih.bold = True

    for inst in general_instructions:
        p_inst = doc.add_paragraph()
        p_inst.paragraph_format.left_indent = Inches(0.15)
        p_inst.paragraph_format.space_after = Pt(1.5)
        r_in = p_inst.add_run(f"• {inst}")
        r_in.font.name = 'Times New Roman'
        r_in.font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 3. Sections Rendering
    for sec in sections_list:
        sec_title = sec.get("section_header", "")
        sec_guidelines = sec.get("guidelines", [])
        
        # Heading 14pt Bold
        p_sec = doc.add_paragraph()
        p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sec.paragraph_format.space_before = Pt(8)
        p_sec.paragraph_format.space_after = Pt(2)
        r_sec = p_sec.add_run(f"\n{sec_title.upper()}")
        r_sec.font.name = 'Times New Roman'
        r_sec.font.size = Pt(14)
        r_sec.bold = True

        for gl in sec_guidelines:
            pg = doc.add_paragraph()
            pg.paragraph_format.left_indent = Inches(0.2)
            pg.paragraph_format.space_after = Pt(1)
            rg = pg.add_run(gl)
            rg.font.name = 'Times New Roman'
            rg.font.size = Pt(11)
            rg.italic = True

        for q in sec.get("questions", []):
            qp = doc.add_paragraph()
            qp.paragraph_format.left_indent = Inches(0.15)
            qp.paragraph_format.space_after = Pt(2)

            q_num = q.get("q_no", "")
            q_header = q.get("instruction_header", "")
            q_text = q.get("question_text", "")
            q_marks = q.get("marks_text", "")
            q_pyq = q.get("pyq_tag", "")

            if q_header:
                r_qh = qp.add_run(f"{q_num} {q_header}  {q_marks}\n")
                r_qh.font.name = 'Times New Roman'
                r_qh.font.size = Pt(12)
                r_qh.bold = True

            if q_text:
                r_qn = qp.add_run(f"{q_num} ")
                r_qn.font.name = 'Times New Roman'
                r_qn.font.size = Pt(12)
                r_qn.bold = True
                
                r_qt = qp.add_run(q_text)
                r_qt.font.name = 'Times New Roman'
                r_qt.font.size = Pt(12)
                
                if q_pyq:
                    r_tag = qp.add_run(f"  [{q_pyq}]")
                    r_tag.font.name = 'Times New Roman'
                    r_tag.font.size = Pt(11)
                    r_tag.bold = True
                    r_tag.font.color.rgb = RGBColor(160, 0, 0)
                
                if q_marks and not q_header:
                    r_qm = qp.add_run(f"   {q_marks}")
                    r_qm.font.name = 'Times New Roman'
                    r_qm.font.size = Pt(12)
                    r_qm.bold = True

            # Options
            if q.get("options") and isinstance(q["options"], dict):
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.left_indent = Inches(0.35)
                opt_p.paragraph_format.space_after = Pt(2)
                opt_line = "      ".join([f"({k}) {v}" for k, v in q["options"].items()])
                r_opt = opt_p.add_run(opt_line)
                r_opt.font.name = 'Times New Roman'
                r_opt.font.size = Pt(12)

            # Sub-items
            sub_items = q.get("sub_items", [])
            for s_idx, sub in enumerate(sub_items, 1):
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.left_indent = Inches(0.35)
                sub_p.paragraph_format.space_after = Pt(1.5)
                
                r_snum = sub_p.add_run(f"{s_idx}. {sub.get('text', '')}")
                r_snum.font.name = 'Times New Roman'
                r_snum.font.size = Pt(12)
                
                sub_pyq = sub.get("pyq_tag", "")
                if sub_pyq:
                    r_stag = sub_p.add_run(f"  [{sub_pyq}]")
                    r_stag.font.name = 'Times New Roman'
                    r_stag.font.size = Pt(11)
                    r_stag.bold = True
                    r_stag.font.color.rgb = RGBColor(160, 0, 0)

                opts = sub.get("options")
                if opts and isinstance(opts, dict):
                    opt_line = "      ".join([f"({k}) {v}" for k, v in opts.items()])
                    r_subopt = sub_p.add_run(f"\n    {opt_line}")
                    r_subopt.font.name = 'Times New Roman'
                    r_subopt.font.size = Pt(12)

    # 4. Save to Disk
    output_dir = "generated_papers"
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    sanitized_sub = subject.replace(" ", "_").replace("/", "_")
    sanitized_class = class_level.replace(" ", "_")
    filename = f"{sanitized_class}_{sanitized_sub}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)
    return filepath, filename